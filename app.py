
import streamlit as st
import PyPDF2
import nltk
from collections import Counter
from docx import Document
import difflib  # For calculating similarity in plagiarism check
from dotenv import load_dotenv
load_dotenv()
import base64
import os
from PIL import Image
import pdf2image
import google.generativeai as genai
from io import BytesIO
from fpdf import FPDF
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io


# Set Streamlit page config at the top
st.set_page_config(page_title="GLA ATS System", page_icon=":guardsman:")

nltk.download('punkt')
nltk.download('stopwords')
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize

# Configure Google Generative AI
genai.configure(api_key=("AIzaSyBPDNB9oDlVpJlTdEkEnc7vWv_CsAZiVQ0"))

def get_gemini_response(resume_text, job_desc_text, prompt):
    """Fetches a response from Gemini API."""
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Combine inputs into a single text blob
        input_text = f"Resume:\n{resume_text}\n\nJob Description:\n{job_desc_text}\n\nPrompt:\n{prompt}"
        
        response = model.generate_content(input_text)
        return response.text
    except Exception as e:
        st.error(f"Error in Gemini API: {e}")
        return None


def extract_text_from_pdf(pdf_file):
    reader = PyPDF2.PdfReader(pdf_file)
    text = ''
    for page_num in range(len(reader.pages)):
        page = reader.pages[page_num]
        text += page.extract_text()
    return text

def input_pdf_setup(pdf_file):
    return [extract_text_from_pdf(pdf_file)]

def create_resume():
    st.title("Create Your Resume")
    
    # Personal Information
    st.header("Personal Information")
    name = st.text_input("Full Name", key="name")
    email = st.text_input("Email", key="email")
    phone = st.text_input("Phone", key="phone")
    linkedin = st.text_input("LinkedIn URL", key="linkedin")
    
    # Professional Summary
    st.header("Professional Summary")
    summary = st.text_area("Enter your professional summary", key="summary")
    
    # Projects
    st.header("Projects")
    num_projects = st.number_input("Number of Projects", min_value=0, max_value=10, value=1, key="num_projects")
    projects = []
    
    for i in range(num_projects):
        st.subheader(f"Project {i+1}")
        project_name = st.text_input(f"Project Name {i+1}", key=f"project_name_{i}")
        project_duration = st.text_input(f"Duration {i+1} (e.g., Jun 2021 – July 2021)", key=f"project_duration_{i}")
        project_details = st.text_area(f"Project Details {i+1}", key=f"project_details_{i}")
        projects.append({
            'name': project_name,
            'duration': project_duration,
            'details': project_details.split('\n')
        })
    
    # Experience
    st.header("Work Experience")
    num_experiences = st.number_input("Number of Experiences", min_value=0, max_value=10, value=1, key="num_experiences")
    experiences = []
    
    for i in range(num_experiences):
        st.subheader(f"Experience {i+1}")
        company_name = st.text_input(f"Company Name {i+1}", key=f"company_name_{i}")
        duration = st.text_input(f"Duration {i+1}", key=f"exp_duration_{i}")
        exp_details = st.text_area(f"Experience Details {i+1}", key=f"exp_details_{i}")
        experiences.append({
            'company': company_name,
            'duration': duration,
            'details': exp_details.split('\n')
        })
    
    # Education
    st.header("Education")
    num_education = st.number_input("Number of Education Entries", min_value=0, max_value=5, value=1, key="num_education")
    education = []
    
    for i in range(num_education):
        st.subheader(f"Education {i+1}")
        institution = st.text_input(f"Institution Name {i+1}", key=f"institution_{i}")
        degree = st.text_input(f"Degree/Certificate {i+1}", key=f"degree_{i}")
        edu_duration = st.text_input(f"Duration {i+1}", key=f"edu_duration_{i}")
        grade = st.text_input(f"Grade/Score {i+1}", key=f"grade_{i}")
        education.append({
            'institution': institution,
            'degree': degree,
            'duration': edu_duration,
            'grade': grade
        })
    
    # Skills
    st.header("Skills")
    skills = st.text_area("Enter your skills (one per line)", key="skills")
    
    # Certifications
    st.header("Certifications")
    certifications = st.text_area("Enter your certifications (one per line)", key="certifications")
    
    if st.button("Generate Resume", key="generate_resume"):
        doc = Document()
        
        # Name and Contact Info
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        contact_paragraph = doc.add_paragraph()
        contact_info = f"{email} | {phone} | {linkedin}"
        contact_run = contact_paragraph.add_run(contact_info)
        contact_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Professional Summary
        doc.add_heading('PROFESSIONAL SUMMARY', level=1)
        doc.add_paragraph(summary)
        
        # Projects
        if projects:
            doc.add_heading('PROJECTS', level=1)
            for project in projects:
                p = doc.add_paragraph()
                p.add_run(f"{project['name']} ({project['duration']})").bold = True
                for detail in project['details']:
                    if detail.strip():
                        doc.add_paragraph(detail.strip(), style='List Bullet')
        
        # Experience
        if experiences:
            doc.add_heading('EXPERIENCE', level=1)
            for exp in experiences:
                p = doc.add_paragraph()
                p.add_run(f"{exp['company']} ({exp['duration']})").bold = True
                for detail in exp['details']:
                    if detail.strip():
                        doc.add_paragraph(detail.strip(), style='List Bullet')
        
        # Education
        if education:
            doc.add_heading('EDUCATION', level=1)
            for edu in education:
                p = doc.add_paragraph()
                p.add_run(f"{edu['institution']} ({edu['duration']})").bold = True
                doc.add_paragraph(f"{edu['degree']} - {edu['grade']}")
        
        # Skills
        if skills:
            doc.add_heading('SKILLS', level=1)
            skills_list = [skill.strip() for skill in skills.split('\n') if skill.strip()]
            doc.add_paragraph(' • '.join(skills_list))
        
        # Certifications
        if certifications:
            doc.add_heading('CERTIFICATION', level=1)
            cert_list = [cert.strip() for cert in certifications.split('\n') if cert.strip()]
            for cert in cert_list:
                doc.add_paragraph(cert, style='List Bullet')
        
        # Save document to bytes buffer
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        # Create download button
        st.download_button(
            label="Download Resume",
            data=doc_bytes,
            file_name="resume.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            key="download_resume"
        )
# Streamlit UI
st.markdown("<h1 style='text-decoration: underline;'>JobFit AI – Smart Hiring Made Easy</h1>", unsafe_allow_html=True)
st.write("🚀 **AI-Powered Resume Matching**")

# Features
st.write("Uses Google Gemini API for advanced NLP & machine learning.")
st.write("Analyzes resumes based on skills, experience, and job relevance.")

st.write("📊 **Automated Ranking System**")
st.write("Scores and ranks candidates based on suitability.")
st.write("Helps recruiters identify top talent instantly.")

st.write("⚡ **Seamless & Efficient Hiring**")
st.write("Streamlit-powered UI for an interactive experience.")
st.write("Reduces manual screening and shortens hiring time.")

st.write("📂 **ATS Integration & Data-Driven Insights**")
st.write("Easily integrates with Applicant Tracking Systems (ATS).")
st.write("Provides real-time analytics for better hiring decisions.")

# Useful Links
st.write("🔗 **Resources**")
st.write("• [Streamlit](https://streamlit.io/)")
st.write("• [Gemini Pro](https://deepmind.google/technologies/gemini/#introduction)")
st.write("• [Makersuite API Key](https://makersuite.google.com/)")
# Sidebar for input
st.sidebar.header("Upload Your Job Description")
job_desc_file = st.sidebar.file_uploader("Upload Job Description (PDF)", type="pdf")


#Prompts
input_prompt1 = """
 You are an experienced Technical Human Resource Manager. Your task is to review the provided resume against the job description.
Please share your professional evaluation on whether the candidate's profile aligns with the role. Highlight the strengths and weaknesses
of the applicant in relation to the specified job requirements.
"""

input_prompt3 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of ATS functionality.
Your task is to evaluate the resume_pdf_content against the provided pdf_content and provide a match percentage.
The output should be a numerical percentage value only, without any additional text or symbols (e.g., 75).
"""


input_prompt4 = """
You are an skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality,
your task is to evaluate the resume against the provided job description. give me the relevant skills if the resume matches
the job description. The output should come as text containing all relevant skills required for given job description .
"""

input_prompt5 = """
You are an skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality,
your task is to evaluate the resume against the provided job description. give me the non-relevant skills if the resume matches
the job description. The output should come as text containing all non-relevant skills mentioned in resume that are not required for given job description .
"""

input_prompt7 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
Your task is to evaluate the resume against the provided job description and return only the Relevant Projects, for the given job description.
The output should come as text containing all relevant projects required for given job description.
"""

input_prompt8 = """
You are a skilled ATS (Applicant Tracking System) scanner with a deep understanding of data science and ATS functionality.
Your task is to evaluate the resume against the provided job description and return only the Recommended Skills required that are not available in resume but given in job description.
The output should come as text containing all recommended skills required for given job description.
"""

# If a job description is uploaded
if job_desc_file is not None:
    op = st.sidebar.selectbox("Resume:", ["Choose an option", "Yes, I have", "No, I have to create."])
    pdf_content = input_pdf_setup(job_desc_file)
    job_desc_text = pdf_content[0]

    # Call the API with the prompts
    if op == "Yes, I have":
        st.subheader("Your Resume")
        resume_file = st.file_uploader("Upload Your Resume (PDF)", type="pdf")

        if resume_file is not None:
            opt = st.sidebar.selectbox("Available Options", ["Choose an option","Percentage match" ,"Show Relevant Skills", "Non-relevant Skills", "Recommended Skills","Relevant Projects","Tell Me About the Resume"])
            resume_pdf_content = input_pdf_setup(resume_file)
            resume_text = resume_pdf_content[0]

            # Get match percentage
            if opt == "Percentage match":
              response = get_gemini_response(input_prompt3, resume_pdf_content, job_desc_text[0])
              # Display the percentage as a progress bar
              st.subheader("Percentage Match")
              st.progress(int(response))
              st.write(f"Match: {response}%")

            # Get relevant skills
            if opt == "Show Relevant Skills":
              relevant_skills = get_gemini_response(resume_text, pdf_content, input_prompt4)
              st.write("Relevant Skills:")
              st.write(relevant_skills)

            # Get non-relevant skills
            if opt == "Non-relevant Skills":
              non_relevant_skills = get_gemini_response(resume_text, pdf_content, input_prompt5)
              st.write("Non-Relevant Skills:")
              st.write(non_relevant_skills)

            # Get relevant projects
            if opt == "Relevant Projects":
              relevant_projects = get_gemini_response(resume_text, pdf_content, input_prompt7)
              st.write("Relevant Projects:")
              st.write(relevant_projects)

            # Get recommended skills
            if opt == "Recommended Skills":
              recommended_skills = get_gemini_response(resume_text, pdf_content, input_prompt8)
              st.write("Recommended Skills:")
              st.write(recommended_skills)

            if opt == "Tell Me About the Resume":
              st.subheader("Detailed Evaluation of Resume")
              evaluation_response = get_gemini_response(resume_pdf_content, pdf_content, input_prompt1)
              if evaluation_response:
                  st.write(evaluation_response)

    if op == "No, I have to create.":
        create_resume()
